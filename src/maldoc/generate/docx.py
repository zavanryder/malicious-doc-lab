"""DOCX document generation with python-docx."""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

from maldoc.attacks.base import AttackResult


def generate_docx(attack_result: AttackResult, title: str, output_path: Path) -> Path:
    """Generate a DOCX containing the attack payload."""
    doc = Document()

    # Metadata injection
    if attack_result.metadata:
        doc.core_properties.author = attack_result.metadata.get("author", "")
        doc.core_properties.subject = attack_result.metadata.get("subject", "")
        doc.core_properties.keywords = attack_result.metadata.get("keywords", "")

    # Title
    doc.add_heading(title, level=1)

    # Visible body
    doc.add_paragraph(attack_result.visible_content)

    # White-on-white: add hidden paragraph
    if attack_result.technique == "white_on_white":
        p = doc.add_paragraph()
        run = p.add_run(attack_result.hidden_content)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(1)

    # Font-size-zero hidden text
    hints = attack_result.format_hints or {}
    if hints.get("font_size_zero_text"):
        p = doc.add_paragraph()
        run = p.add_run(hints["font_size_zero_text"])
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(1)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
